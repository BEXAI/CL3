"""
Generate the SYH Phase 2 Scope of Engagement document.
Palatino Linotype, all black, no dashes (use 'to' or commas), single .docx.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BLACK = RGBColor(0, 0, 0)
FONT_NAME = "Palatino Linotype"
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "SYH_Phase2_Scope_of_Engagement.docx")


def set_font(run, size=11, bold=False, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.bold = bold
    run.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = r.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_paragraph(doc, text, size=11, bold=False, italic=False, alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_mixed_paragraph(doc, segments, alignment=None, space_after=6, space_before=0):
    """segments: list of (text, size, bold, italic) tuples"""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    for text, size, bold, italic in segments:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading_styled(doc, text, size=14, space_before=18, space_after=6):
    return add_paragraph(doc, text, size=size, bold=True, space_before=space_before, space_after=space_after)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=10, bold=True)
        # Light gray background for header
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "E6E6E6"
        })
        shading.append(shd)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return table


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # --- TITLE ---
    add_paragraph(doc, "", size=11, space_after=36)  # top spacer
    add_paragraph(doc, "Superyacht Holdings",
                  size=26, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc, "Phase 2: Scope of Engagement",
                  size=18, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc, "Revenue Engine, Member Operations, and Fleet Economics",
                  size=12, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    add_paragraph(doc, "Prepared for SYH Club Leadership",
                  size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "June 2026",
                  size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Confidential",
                  size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)

    # --- PAGE BREAK ---
    doc.add_page_break()

    # ===========================
    # PHASE 1 RECAP
    # ===========================
    add_heading_styled(doc, "Phase 1 Recap", size=16, space_before=6)

    add_paragraph(doc, (
        "Phase 1 delivered the prospect intelligence asset. A multi-pass agentic enrichment pipeline "
        "processed 25,003 UHNW contact records from the All WE Contacts dataset, resolved each "
        "record against LinkedIn via a five-pass Google search strategy, and scored every match "
        "with a confidence engine that weighs name verification, company alignment, geographic "
        "signals, and URL slug quality."
    ))

    add_heading_styled(doc, "Consolidated Results", size=12, space_before=12)

    add_table(doc,
              ["Metric", "Value"],
              [
                  ["Total records processed", "25,003"],
                  ["High confidence matches (score 80 or above)", "21,864 (87.4%)"],
                  ["Low confidence with URL (below 80)", "2,732 (10.9%)"],
                  ["No LinkedIn URL found", "407 (1.6%)"],
                  ["Total enriched columns added", "6 (URL, Title, Description, Score, Signals, Pass)"],
              ],
              col_widths=[3.5, 2.5])

    add_heading_styled(doc, "Propensity Scoring (Phase 1 Extension)", size=12, space_before=12)

    add_paragraph(doc, (
        "The enriched dataset was then scored for founding member propensity across six dimensions: "
        "title seniority (0 to 25), wealth composite (0 to 30), real estate and aviation footprint "
        "(0 to 15), philanthropy and influence (0 to 15), business ownership (0 to 10), and LinkedIn "
        "profile strength (0 to 5). The resulting 100 point composite score segments the full 25,003 "
        "records into actionable tiers."
    ))

    add_table(doc,
              ["Tier", "Score Range", "Count", "Share"],
              [
                  ["Platinum", "75 to 100", "668", "2.7%"],
                  ["Gold", "60 to 74", "3,800", "15.2%"],
                  ["Silver", "45 to 59", "8,456", "33.8%"],
                  ["Bronze", "30 to 44", "11,233", "44.9%"],
                  ["Prospect", "0 to 29", "846", "3.4%"],
              ],
              col_widths=[1.5, 1.5, 1.5, 1.5])

    add_paragraph(doc, (
        "The 668 Platinum and 3,800 Gold tier records represent the immediate founding member "
        "pipeline. Phase 2 activates this asset into a closing and operating engine."
    ), space_before=6)

    # --- PAGE BREAK ---
    doc.add_page_break()

    # ===========================
    # PHASE 2 OVERVIEW
    # ===========================
    add_heading_styled(doc, "Phase 2 Overview", size=16, space_before=6)

    add_paragraph(doc, (
        "Phase 1 built the asset. Phase 2 activates it into a revenue and member operations engine. "
        "The scope is organized into three layers: the conversion engine (Items 1 through 4), the "
        "operating backbone (Items 5 through 8), and the visibility and trust layer (Items 9 and 10). "
        "Each deliverable runs on the same finance-first, build-real-tools positioning that informed "
        "Phase 1."
    ))

    # Architecture summary
    add_heading_styled(doc, "Scope Architecture", size=12, space_before=12)
    add_table(doc,
              ["Layer", "Items", "Function"],
              [
                  ["Conversion Engine", "1 through 4", "Turn the enriched prospect list into closed founding members"],
                  ["Operating Backbone", "5 through 8", "Run the yacht, the credits, the billing, and the member experience"],
                  ["Visibility and Trust", "9 and 10", "Give leadership a single source of truth and protect the data"],
              ],
              col_widths=[1.8, 1.2, 3.5])

    # --- PAGE BREAK ---
    doc.add_page_break()

    # ===========================
    # CONVERSION ENGINE (ITEMS 1-4)
    # ===========================
    add_heading_styled(doc, "Conversion Engine", size=16, space_before=6)

    # --- ITEM 1 ---
    add_heading_styled(doc, "1. Prospect Qualification and Propensity Scoring", size=13, space_before=14)
    add_paragraph(doc, (
        "Extend the Phase 1 confidence scorer from 'is this the right person' into 'is this the "
        "right member.' Layer wealth and fit signals onto the enriched profiles and rank the 25,003 "
        "records into a prioritized founding member tier list. This is the natural continuation of "
        "the work already shipped."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "Title seniority scoring: C-suite, founder/owner, board, investor, president/EVP weighted by relevance to UHNW membership",
        "Wealth composite signals: board seats, founder/exit history, estimated net worth bands from source data",
        "Real estate and aviation footprint: property counts, value bands, and aircraft registry signals as lifestyle proxies",
        "Philanthropy and influence: foundation affiliations, donor history, and public board service",
        "Business ownership: active operating companies, holding structures, and family office indicators",
        "Tier classification into Platinum (75 to 100), Gold (60 to 74), Silver (45 to 59), Bronze (30 to 44), Prospect (0 to 29)",
        "Prioritized outreach list ranked by composite score within each tier",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "Scored CSV and pipeline-ready extract with 158+ columns. Tier summary analytics. "
        "Scoring methodology document for client review and regulatory defensibility."
    ), size=10)

    # --- ITEM 2 ---
    add_heading_styled(doc, "2. Agentic Outreach and Sequencing", size=13, space_before=14)
    add_paragraph(doc, (
        "A per-prospect engine that drafts personalized multi-channel outreach off the enriched data, "
        "schedules touches, and tracks responses. The comparable benchmark in wealth management is "
        "meaningful: agentic prospecting reduces advisor time on manual prospecting by 40 to 50% and "
        "lifts net new AUM by 30 to 40%."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "LinkedIn outreach: personalized connection request and follow-up message drafts keyed to prospect tier, title, and known interests",
        "Email sequences: three to five touch cadence per prospect with merge fields from enriched columns (company, title, location, wealth signals)",
        "Warm intro mapping: cross-reference the 25,003 against existing SYH member networks to identify shared board seats, alumni, or philanthropic affiliations",
        "Response tracking: log opens, replies, and meeting bookings per prospect with cadence pause on engagement",
        "Agent-generated content: each draft grounded in the prospect's enriched profile, not generic templates",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "Outreach engine integrated with the scored prospect list. Sequence templates by tier. "
        "Warm intro graph. Response dashboard."
    ), size=10)

    # --- ITEM 3 ---
    add_heading_styled(doc, "3. CRM and Pipeline System of Record", size=13, space_before=14)
    add_paragraph(doc, (
        "Move SYH off raw CSVs into a real CRM seeded with the enriched columns and scores, with "
        "founding member stages and a live pipeline view. Right now the intelligence lives in a file. "
        "Phase 2 makes it a system the sales team runs on."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "CRM selection and provisioning: evaluate fit against SYH's sales motion (HubSpot, Salesforce, or purpose-built depending on team size and budget)",
        "Data migration: seed the CRM with all 158 enriched columns, tier assignments, and propensity scores from the Phase 1 master dataset",
        "Founding member pipeline stages: Application, Invitation, Discovery Call, Due Diligence, Close, Onboarded",
        "Pipeline views: Kanban board by stage, filterable by tier, geography, and wealth band",
        "Reporting: weekly pipeline value, conversion rate by tier, and stage velocity",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "Configured CRM with seeded data, pipeline stages, and reporting views. "
        "Data dictionary mapping enriched columns to CRM fields."
    ), size=10)

    # --- ITEM 4 ---
    add_heading_styled(doc, "4. KYC, AML, and Source of Wealth Onboarding", size=13, space_before=14)
    add_paragraph(doc, (
        "Issuing equity to UHNW members is a regulated transaction. Build an agentic onboarding flow "
        "that verifies identity, screens sanctions and PEP lists, and documents source of wealth before "
        "equity is issued. An estimated 70% of new account onboarding is projected to be fully automated "
        "by 2026, making this table stakes for any equity-issuing membership structure. The compliance "
        "scope aligns with SEC, FINRA, and NIST frameworks."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "Identity verification: document upload, liveness check, and name/address match against source records",
        "Sanctions and PEP screening: automated checks against OFAC, EU, and UN consolidated lists plus politically exposed persons databases",
        "Source of wealth documentation: structured intake for employment history, business ownership, investment proceeds, inheritance, and real estate",
        "Ongoing monitoring: periodic re-screening triggers for sanctions list updates and adverse media",
        "Audit trail: timestamped record of every verification step, decision, and override for regulatory examination",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "Onboarding workflow with identity, sanctions, PEP, and source of wealth modules. "
        "Compliance documentation package. Audit log schema."
    ), size=10)

    # --- PAGE BREAK ---
    doc.add_page_break()

    # ===========================
    # OPERATING BACKBONE (ITEMS 5-8)
    # ===========================
    add_heading_styled(doc, "Operating Backbone", size=16, space_before=6)

    # --- ITEM 5 ---
    add_heading_styled(doc, "5. Fractional Ownership Financial Model", size=13, space_before=14)
    add_paragraph(doc, (
        "Build the unit economics and financial model for SYH's fractional yacht product. Vessel "
        "acquisition plus operating expenditure spread across 78 owners, membership pricing, credit "
        "system valuation, breakeven per hull, and fleet expansion scenarios in FAST Standard."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "Three-statement or cash flow model: revenue (membership dues, assessments, capital calls), operating costs (crew, fuel, port fees, maintenance, insurance), and capital structure",
        "Unit economics per member: acquisition cost (equity share), annual operating share, credit value, and effective cost per night",
        "Breakeven analysis: minimum membership count per hull, utilization thresholds, and dues sensitivity",
        "Scenario modeling in FAST Standard: Base, Best, Worst, and Stress cases with clearly stated assumptions",
        "Fleet expansion economics: marginal cost and revenue per additional hull, cross-fleet credit dilution, and capital call sizing",
        "Credit system valuation: economic value of a credit unit, stateroom size adjustment factors, and reciprocity exchange rates across multiple vessels",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "FAST Standard financial model in Excel. Scenario dashboard. "
        "Assumptions document and sensitivity tables."
    ), size=10)

    # --- ITEM 6 ---
    add_heading_styled(doc, "6. Booking and Credit Reciprocity Engine", size=13, space_before=14)
    add_paragraph(doc, (
        "The credit system adjusts for stateroom size, enables reciprocity across the fleet as SYH "
        "expands, and lets members trade time onboard. This is a non-trivial allocation and booking "
        "system that underpins the member experience."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "Credit allocation engine: annual credit issuance per member based on equity share and stateroom tier",
        "Stateroom sizing rules: credit multipliers by cabin class (owner's suite, VIP, standard) with seasonal adjustments for peak periods (Monaco Grand Prix, St. Barths New Year's)",
        "Booking system: member-facing interface for browsing availability, reserving nights, and managing waitlists across Mediterranean summer and Caribbean winter itineraries",
        "Fleet reciprocity: credit portability across vessels as SYH adds hulls, with exchange rate logic that preserves per-member value",
        "Member-to-member trading: marketplace for buying, selling, or gifting credits between members with transaction logging",
        "Inventory management: real-time stateroom availability, blackout date handling, and overbooking prevention",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "Rules engine specification. Member-facing booking interface. Credit ledger with "
        "transaction history. Fleet reciprocity rate tables."
    ), size=10)

    # --- ITEM 7 ---
    add_heading_styled(doc, "7. Member Concierge Assistant", size=13, space_before=14)
    add_paragraph(doc, (
        "An on-brand AI concierge for itineraries, marquee events, onboard requests, and stored "
        "preferences. This is a differentiated member experience layer that most clubs cannot build, "
        "and it draws directly on demonstrated experience shipping a working AI assistant."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "Itinerary planning: Mediterranean and Caribbean route recommendations with port-specific dining, excursion, and event suggestions",
        "Marquee event concierge: Cannes Film Festival, Monaco Grand Prix, St. Barths New Year's, and similar calendar events with reservation and logistics coordination",
        "Onboard requests: dining preferences, spa bookings, water toy reservations, and special occasion arrangements routed to crew",
        "Member preference memory: stored dietary restrictions, cabin temperature, pillow type, beverage preferences, and past request history",
        "Multi-channel access: web, mobile, and messaging interface with natural language interaction",
        "Brand voice: all interactions reflect SYH's tone, exclusivity, and service standards",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "Deployed AI concierge with preference engine, event knowledge base, and crew routing. "
        "Brand voice guidelines document."
    ), size=10)

    # --- ITEM 8 ---
    add_heading_styled(doc, "8. Dues, Billing, and Payments Automation", size=13, space_before=14)
    add_paragraph(doc, (
        "Recurring membership dues, operating expense assessments, capital calls for fleet expansion, "
        "and ACH/wire reconciliation. Autonomous reconciliation and close management is one of the "
        "highest-return finance automations available today, and it ties cleanly to existing QuickBooks "
        "and finance stack integrations."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "Recurring billing: monthly or quarterly dues invoicing with automated ACH/wire collection",
        "Operating expense assessments: pro-rata opex charges based on credit usage and stateroom tier",
        "Capital calls: fleet expansion funding requests with waterfall allocation logic, notice periods, and payment tracking",
        "Payment reconciliation: automated matching of incoming payments against invoices with exception flagging",
        "Accounting integration: sync to QuickBooks or equivalent for general ledger posting, revenue recognition, and financial reporting",
        "Member portal: self-service billing history, payment method management, and statement downloads",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "Billing engine with recurring, assessment, and capital call modules. "
        "Reconciliation automation. Accounting system integration. Member billing portal."
    ), size=10)

    # --- PAGE BREAK ---
    doc.add_page_break()

    # ===========================
    # VISIBILITY AND TRUST (ITEMS 9-10)
    # ===========================
    add_heading_styled(doc, "Visibility and Trust", size=16, space_before=6)

    # --- ITEM 9 ---
    add_heading_styled(doc, "9. Executive Dashboard", size=13, space_before=14)
    add_paragraph(doc, (
        "One source of truth for leadership: pipeline value, member acquisition cost, utilization "
        "per yacht, dues collection and DSO, and fleet P&L. This is the layer that lets SYH "
        "leadership see the whole funnel and the fleet economics in real time."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "Pipeline dashboard: total pipeline value by stage, conversion rates by tier, stage velocity, and forecast",
        "Member acquisition metrics: cost per member acquired, time to close by tier, and channel attribution",
        "Fleet utilization: nights booked vs. available by vessel, stateroom class, and season with occupancy trends",
        "Financial health: dues collection rate, days sales outstanding, operating margin per vessel, and fleet-level P&L",
        "Credit economy: credits issued vs. redeemed, trading volume, and credit utilization rate",
        "Real-time refresh: dashboard pulls from CRM, billing, and booking systems with automated data pipeline",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "Live executive dashboard with pipeline, utilization, financial, and credit economy views. "
        "Automated data refresh pipeline. KPI definitions document."
    ), size=10)

    # --- ITEM 10 ---
    add_heading_styled(doc, "10. Data Governance, Security, and Integrity", size=13, space_before=14)
    add_paragraph(doc, (
        "A UHNW member list is among the most sensitive data assets that exists, so privacy and "
        "access control are not optional. This workstream establishes the controls, audit mechanisms, "
        "and data quality standards that underpin every other Phase 2 deliverable."
    ))
    add_heading_styled(doc, "Scope", size=11, space_before=8)
    for item in [
        "Role-based access control: define roles (sales, operations, finance, leadership, compliance) with granular permissions per dataset and system",
        "Audit logging: timestamped record of every data access, modification, agent action, and export with user attribution",
        "Explainable agent actions: every automated decision (score, outreach draft, KYC determination) includes a confidence signal breakdown and reasoning trace",
        "PII handling: encryption at rest and in transit, data minimization principles, retention policies, and member consent management",
        "Secrets management: API keys, tokens, and credentials stored in environment variables or a secrets manager, never in source code",
        "Verified data principle: all member-facing data sourced from confirmed records only, with unverified fields flagged and excluded from operational use",
        "Incident response: data breach notification plan, access revocation procedures, and forensic audit capability",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)

    add_heading_styled(doc, "Deliverable", size=11, space_before=6)
    add_paragraph(doc, (
        "Data governance policy. RBAC matrix. Audit log schema and retention policy. "
        "Secrets management implementation. PII handling procedures."
    ), size=10)

    # --- PAGE BREAK ---
    doc.add_page_break()

    # ===========================
    # DATA INTEGRITY NOTES
    # ===========================
    add_heading_styled(doc, "Data Integrity Notes", size=16, space_before=6)

    add_heading_styled(doc, "Match Rate Methodology", size=12, space_before=12)
    add_paragraph(doc, (
        "The headline match rate of 87.4% represents the final consolidated result across all "
        "25,003 records after a five-pass multi-query strategy with enhanced confidence scoring. "
        "21,864 records achieved a confidence score of 80 or above. Intermediate batch summaries "
        "produced during processing (including an early single-pass result of 10% on the first "
        "5,000 records before reformulation) reflect partial pipeline state and are not representative "
        "of final yield. All client-facing materials reference the consolidated figure only."
    ))

    add_heading_styled(doc, "Batch Consistency", size=12, space_before=12)
    add_table(doc,
              ["Batch", "Records", "High Confidence", "Rate"],
              [
                  ["1", "1 to 5,000", "4,358", "87.2%"],
                  ["2", "5,001 to 10,000", "4,354", "87.1%"],
                  ["3", "10,001 to 15,000", "4,416", "88.3%"],
                  ["4", "15,001 to 20,000", "4,390", "87.8%"],
                  ["5", "20,001 to 25,003", "4,346", "86.9%"],
                  ["Consolidated", "1 to 25,003", "21,864", "87.4%"],
              ],
              col_widths=[1.5, 1.5, 1.5, 1.5])

    add_paragraph(doc, (
        "Per-batch rates range from 86.9% to 88.3%, confirming consistent performance across "
        "the full dataset with no batch-specific anomalies."
    ))

    add_heading_styled(doc, "Security Remediation", size=12, space_before=12)
    add_paragraph(doc, (
        "Phase 1 processing scripts contain a hardcoded Apify API token in seven Python files. "
        "This token must be rotated in the Apify console and replaced with an environment variable "
        "reference before any code is shared with the client or committed to a shared repository. "
        "This remediation falls under Item 10 (Data Governance, Security, and Integrity) and will "
        "be completed as a prerequisite action before Phase 2 delivery begins."
    ))

    # --- PAGE BREAK ---
    doc.add_page_break()

    # ===========================
    # ENGAGEMENT STRUCTURE
    # ===========================
    add_heading_styled(doc, "Engagement Structure", size=16, space_before=6)

    add_heading_styled(doc, "Phasing", size=12, space_before=12)
    add_paragraph(doc, (
        "Items are ordered by dependency and value realization. The conversion engine (Items 1 "
        "through 4) can begin immediately because it builds directly on Phase 1 deliverables. "
        "The operating backbone (Items 5 through 8) can run in parallel once founding member "
        "pipeline stages are defined. The visibility and trust layer (Items 9 and 10) spans the "
        "full engagement as a cross-cutting workstream."
    ))

    add_table(doc,
              ["Priority", "Items", "Dependency"],
              [
                  ["Immediate", "1. Propensity Scoring", "Phase 1 dataset (delivered)"],
                  ["Immediate", "3. CRM and Pipeline", "Phase 1 dataset (delivered)"],
                  ["Immediate", "10. Data Governance", "None (cross-cutting)"],
                  ["Fast follow", "2. Agentic Outreach", "Requires Item 1 scoring and Item 3 CRM"],
                  ["Fast follow", "4. KYC/AML Onboarding", "Requires Item 3 pipeline stages"],
                  ["Fast follow", "5. Financial Model", "Requires SYH vessel and pricing data"],
                  ["Parallel", "6. Booking and Credits", "Requires Item 5 credit valuation"],
                  ["Parallel", "7. Member Concierge", "Requires Item 6 booking system"],
                  ["Parallel", "8. Dues and Billing", "Requires Item 5 pricing model"],
                  ["Continuous", "9. Executive Dashboard", "Aggregates data from Items 1 through 8"],
              ],
              col_widths=[1.2, 2.3, 3.0])

    add_heading_styled(doc, "Positioning", size=12, space_before=12)
    add_paragraph(doc, (
        "Phase 1 built the asset. Phase 2 activates it into a revenue and member operations engine. "
        "Every deliverable is grounded in the same discipline: build real tools that the team runs "
        "on, model the economics so leadership can see what is true, and treat data integrity as "
        "a first-class requirement rather than an afterthought."
    ))

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
